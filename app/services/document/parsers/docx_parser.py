"""Microsoft Word document parser implementation."""

import io
from typing import Any

import docx

from app.services.document.parser_base import (
    DocumentParser,
    ParserRegistry,
    ParsingResult,
)


class DocxParser(DocumentParser):
    """Parser for Microsoft Word (.docx) documents."""

    @classmethod
    def supported_mimetypes(cls) -> list[str]:
        """Return list of supported MIME types."""
        return [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword',  # For older .doc files (limited support)
        ]

    async def parse(self, content: bytes, filename: str | None = None) -> ParsingResult:
        """Parse Word document content."""
        try:
            doc = docx.Document(io.BytesIO(content))

            # Extract full text
            text = self._extract_text(doc)

            # Extract document properties/metadata
            metadata = self._extract_metadata(doc)

            # Extract document structure
            structure = self._extract_structure(doc)

            return ParsingResult(
                text=text,
                metadata=metadata,
                pages=structure.get('estimated_pages', 1),
                structure=structure,
            )
        except Exception as e:
            return ParsingResult(text='', error=f'Word document parsing failed: {e!s}')

    def _extract_text(self, doc: docx.Document) -> str:
        """Extract all text from the document."""
        try:
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            return f'Error extracting text: {e!s}'

    def _extract_metadata(self, doc: docx.Document) -> dict[str, Any]:
        """Extract document metadata."""
        core_props = doc.core_properties

        metadata = {
            'title': core_props.title or '',
            'author': core_props.author or '',
            'subject': core_props.subject or '',
            'keywords': core_props.keywords or '',
            'created': (core_props.created and str(core_props.created)) or '',
            'modified': (core_props.modified and str(core_props.modified)) or '',
            'last_modified_by': core_props.last_modified_by or '',
            'category': core_props.category or '',
            'language': core_props.language or '',
        }

        return metadata

    def _extract_structure(self, doc: docx.Document) -> dict[str, Any]:
        """Extract document structure."""
        # Extract headings by analyzing paragraphs with heading styles
        headings = []
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                headings.append(
                    {
                        'level': int(paragraph.style.name.replace('Heading', ''))
                        if paragraph.style.name != 'Heading'
                        else 1,
                        'text': paragraph.text,
                    }
                )

        # Count tables and images
        tables_count = len(doc.tables)

        # Estimate pages - improved algorithm
        # Count total characters and use a more realistic estimate
        total_chars = sum(len(paragraph.text) for paragraph in doc.paragraphs)

        # A4 page with standard margins holds about 3000 characters per page
        estimated_pages = max(1, round(total_chars / 3000))

        # Adjust for tables (each table adds approximately 0.5 pages)
        estimated_pages += tables_count * 0.5

        # Round up to whole number of pages
        estimated_pages = int(estimated_pages + 0.5)

        return {
            'headings': headings,
            'tables_count': tables_count,
            'paragraphs_count': len(doc.paragraphs),
            'character_count': total_chars,
            'estimated_pages': estimated_pages,
        }


# Register the parser
ParserRegistry.register(DocxParser)
